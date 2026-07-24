from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REFERENCE = Path(r"C:\Users\ABRAR\OneDrive\Desktop\Lindafashion-Automation\tmp\labels\LINDA-LABEL-reference.docx")
OUTPUT = Path(r"C:\Users\ABRAR\OneDrive\Desktop\Lindafashion-Automation\tmp\labels\LABEL-15914-sample.docx")

ADDRESS = [
    "AMAD LOGISTICS",
    "IMPORTACIONES MUNDO OFERTAS-SI CARG",
    "8274 N.W. 14TH STREET",
    "305 888-0344",
    "DORAL, FL 33126",
    "C/O BELLAS INTERNACIONAL LLC.",
]


def borders(cell, **edges):
    props = cell._tc.get_or_add_tcPr()
    node = props.first_child_found_in("w:tcBorders")
    if node is None:
        node = OxmlElement("w:tcBorders")
        props.append(node)
    for edge, attrs in edges.items():
        tag = f"w:{edge}"
        element = node.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            node.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_margin(cell, top=55, start=55, bottom=55, end=55):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def compact(paragraph, before=0, after=0, line=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


with ZipFile(REFERENCE) as archive:
    logo = BytesIO(archive.read("word/media/image1.jpeg"))

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.05)
section.bottom_margin = Inches(0.05)
section.left_margin = Inches(0.05)
section.right_margin = Inches(0.05)
section.header_distance = Inches(0)
section.footer_distance = Inches(0)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(0)

outer = doc.add_table(rows=2, cols=2)
outer.autofit = False
outer.allow_autofit = False
outer.columns[0].width = Inches(4.18)
outer.columns[1].width = Inches(4.18)

for row in outer.rows:
    row.height = Inches(5.42)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for cell in row.cells:
        cell.width = Inches(4.18)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margin(cell)
        borders(cell,
                top={"val": "single", "sz": "6", "color": "000000"},
                bottom={"val": "single", "sz": "6", "color": "000000"},
                start={"val": "single", "sz": "6", "color": "000000"},
                end={"val": "single", "sz": "6", "color": "000000"})

        p0 = cell.paragraphs[0]
        compact(p0)
        header = cell.add_table(rows=1, cols=2)
        header.autofit = False
        header.columns[0].width = Inches(1.30)
        header.columns[1].width = Inches(2.70)
        hp = header.cell(0, 0).paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact(hp)
        logo.seek(0)
        hp.add_run().add_picture(logo, width=Inches(1.15))

        company = header.cell(0, 1).paragraphs[0]
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact(company, line=0.92)
        lines = [
            ("LINDA FASHION ACCESSORIES CORP", 10.5),
            ("2195 ELIZABETH AVE  1ST FL S-BLDG,", 9.5),
            ("RAHWAY, NJ 07065", 9.5),
            ("TEL: 732-669-7263  FAX: 732-540-8475", 9.0),
            ("Website: www.lindafashionny.com", 7.5),
            ("Email: info@lindafashionny.com", 7.5),
        ]
        for index, (text, size) in enumerate(lines):
            run = company.add_run(text)
            run.bold = True
            run.font.size = Pt(size)
            if index != len(lines) - 1:
                run.add_break()

        divider = cell.add_paragraph("--------------------------------------------------------------")
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact(divider, before=1, after=2)
        divider.runs[0].bold = True
        divider.runs[0].font.size = Pt(9)

        ship = cell.add_paragraph("SHIP TO:")
        ship.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact(ship, after=2)
        ship.runs[0].bold = True
        ship.runs[0].font.name = "Times New Roman"
        ship.runs[0].font.size = Pt(21)

        address = cell.add_paragraph()
        address.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact(address, line=1.0)
        for index, text in enumerate(ADDRESS):
            run = address.add_run(text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(14 if index == 1 else 16)
            if index != len(ADDRESS) - 1:
                run.add_break()

        ucl = cell.add_paragraph("UCL")
        ucl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact(ucl, before=2)
        ucl.runs[0].bold = True
        ucl.runs[0].font.name = "Times New Roman"
        ucl.runs[0].font.size = Pt(20)

doc.save(OUTPUT)
print(OUTPUT)
